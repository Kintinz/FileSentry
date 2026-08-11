#!/usr/bin/env python3
"""Generate the end-user Word guide and the internal build certificate.

The build release intentionally contains only three artifacts:
  * FileSentrySentinel.exe
  * FileSentrySentinel_User_Guide.docx
  * FileSentrySentinel_Exclusive_Build_Certificate.docx

The guide is written for ordinary users rather than developers.  It contains
step-by-step flows and diagrams that explain what the user should click and
what result to expect.
"""

from __future__ import annotations

import argparse
import hashlib
import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PRODUCT = "FileSentry Sentinel"
NAVY = "0B1B33"
NAVY_2 = "132B4D"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
CYAN = "2DB7E5"
INK = "0B2545"
MUTED = "60738B"
LIGHT_BLUE = "E8EEF5"
LIGHT = "F4F6F9"
GREEN = "16794A"
AMBER = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def find_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path(os.environ.get("WINDIR", r"C:\\Windows")) / "Fonts" / ("segoeuib.ttf" if bold else "segoeui.ttf"),
        Path(os.environ.get("WINDIR", r"C:\\Windows")) / "Fonts" / ("arialbd.ttf" if bold else "arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if not current or draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def rounded_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], fill: str, outline: str = CYAN, width: int = 4) -> None:
    draw.rounded_rectangle(xy, radius=24, fill=fill if fill.startswith("#") else f"#{fill}", outline=outline if outline.startswith("#") else f"#{outline}", width=width)


def centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.ImageFont, fill: str = WHITE) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 34)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total = sum(heights) + max(0, len(lines) - 1) * 8
    y = y1 + ((y2 - y1) - total) // 2
    for line, height in zip(lines, heights):
        width = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + ((x2 - x1) - width) // 2, y), line, font=font, fill=fill if fill.startswith("#") else f"#{fill}")
        y += height + 8


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = CYAN, width: int = 8) -> None:
    fill = fill if fill.startswith("#") else f"#{fill}"
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    if abs(end[0] - start[0]) >= abs(end[1] - start[1]):
        direction = 1 if end[0] >= start[0] else -1
        points = [(ex, ey), (ex - direction * 22, ey - 14), (ex - direction * 22, ey + 14)]
    else:
        direction = 1 if end[1] >= start[1] else -1
        points = [(ex, ey), (ex - 14, ey - direction * 22), (ex + 14, ey - direction * 22)]
    draw.polygon(points, fill=fill)


def make_logo(path: Path) -> None:
    image = Image.new("RGBA", (640, 640), (0, 0, 0, 0))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((50, 50, 590, 590), radius=120, fill=f"#{NAVY}", outline=f"#{CYAN}", width=14)
    shield = [(320, 110), (465, 175), (445, 380), (320, 505), (195, 380), (175, 175)]
    draw.polygon(shield, fill=f"#{BLUE}", outline=f"#{CYAN}")
    draw.line([(245, 320), (300, 375), (405, 245)], fill=f"#{WHITE}", width=34, joint="curve")
    image.save(path)


def make_user_flow(path: Path) -> None:
    image = Image.new("RGB", (1800, 980), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    title_font = find_font(42, True)
    box_font = find_font(30, True)
    small_font = find_font(24)
    draw.text((70, 40), "Luồng thao tác an toàn trong FileSentry", font=title_font, fill=f"#{WHITE}")
    boxes = [
        (70, 245, 370, 505, "1. Mở app", "Chọn FileSentry Sentinel"),
        (445, 245, 745, 505, "2. Xác thực", "Nhập mật khẩu trong app"),
        (820, 245, 1120, 505, "3. Chọn quyền", "Thư mục / camera / mic"),
        (1195, 245, 1495, 505, "4. Windows", "Thao tác thật diễn ra"),
    ]
    for x1, y1, x2, y2, label, detail in boxes:
        rounded_box(draw, (x1, y1, x2, y2), f"#{NAVY_2}")
        centered_text(draw, (x1 + 18, y1 + 22, x2 - 18, y1 + 130), label, box_font)
        centered_text(draw, (x1 + 18, y1 + 132, x2 - 18, y2 - 22), detail, small_font, f"#{CYAN}")
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[2] + 20, 375), (right[0] - 20, 375))
    rounded_box(draw, (455, 650, 1345, 835), f"#{GREEN}", outline=f"#{CYAN}")
    centered_text(draw, (480, 675, 1320, 810), "Sau khi xác thực: không nhập lại mật khẩu ở Windows trong 15 phút", box_font)
    draw.text((70, 900), "Hết thời hạn, khóa lại hoặc đổi ngoài ý muốn → FileSentry cảnh báo và yêu cầu xác thực lại.", font=small_font, fill=f"#{WHITE}")
    image.save(path)


def make_protection_cycle(path: Path) -> None:
    image = Image.new("RGB", (1800, 860), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    title_font = find_font(42, True)
    box_font = find_font(28, True)
    small_font = find_font(23)
    draw.text((70, 38), "Chu kỳ bảo vệ của khu vực / thư mục", font=title_font, fill=f"#{WHITE}")
    items = [
        (95, 300, 410, 540, "ĐANG KHÓA", "Không mở / không truy cập"),
        (520, 300, 835, 540, "MỞ BẢO VỆ", "Đã xác thực trong app"),
        (945, 300, 1260, 540, "TẠM DỪNG", "Cho phép trong thời gian ngắn"),
        (1370, 300, 1685, 540, "KHÓA LẠI", "Trở về trạng thái an toàn"),
    ]
    colors = [RED, GREEN, AMBER, BLUE]
    for item, color in zip(items, colors):
        x1, y1, x2, y2, label, detail = item
        rounded_box(draw, (x1, y1, x2, y2), f"#{NAVY_2}", outline=f"#{color}")
        centered_text(draw, (x1 + 12, y1 + 25, x2 - 12, y1 + 105), label, box_font, f"#{color}")
        centered_text(draw, (x1 + 18, y1 + 135, x2 - 18, y2 - 25), detail, small_font, f"#{WHITE}")
    for left, right in zip(items, items[1:]):
        arrow(draw, (left[2] + 18, 420), (right[0] - 18, 420), fill=f"#{CYAN}")
    draw.text((95, 680), "Khuyến nghị: sau khi sử dụng xong, bấm “KHÓA LẠI” thay vì để phiên mở lâu.", font=small_font, fill=f"#{WHITE}")
    image.save(path)


def make_media_flow(path: Path) -> None:
    image = Image.new("RGB", (1800, 930), f"#{NAVY}")
    draw = ImageDraw.Draw(image)
    title_font = find_font(42, True)
    box_font = find_font(28, True)
    small_font = find_font(22)
    draw.text((70, 38), "Luồng mở camera / microphone", font=title_font, fill=f"#{WHITE}")
    boxes = [
        (80, 250, 380, 520, "Chọn Camera / Mic", "Trong màn hình Privacy"),
        (480, 250, 780, 520, "Nhập mật khẩu", "Một cổng xác thực duy nhất"),
        (880, 250, 1180, 520, "Mở phiên", "Thời hạn theo policy"),
        (1280, 250, 1580, 520, "Dùng trực tiếp", "Windows / trình duyệt"),
    ]
    for x1, y1, x2, y2, label, detail in boxes:
        rounded_box(draw, (x1, y1, x2, y2), f"#{NAVY_2}")
        centered_text(draw, (x1 + 15, y1 + 22, x2 - 15, y1 + 125), label, box_font)
        centered_text(draw, (x1 + 15, y1 + 140, x2 - 15, y2 - 20), detail, small_font, f"#{CYAN}")
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[2] + 20, 385), (right[0] - 20, 385))
    rounded_box(draw, (270, 655, 1530, 815), f"#{AMBER}", outline=f"#{CYAN}")
    centered_text(draw, (300, 680, 1500, 790), "Nếu quyền thay đổi ngoài dự kiến hoặc hết thời gian → trạng thái trở về khóa và bạn nhận được thông báo.", box_font)
    draw.text((80, 870), "Lưu ý: Windows có thể còn hiển thị UI riêng; FileSentry là lớp cổng + giám sát, không thay thế nhân hệ điều hành.", font=small_font, fill=f"#{WHITE}")
    image.save(path)


def set_run_font(run, size: float = 11, color: str = INK, bold: bool | None = None, italic: bool | None = None, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "C7D3E0", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def configure_document(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, BLUE_DARK, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run(PRODUCT)
    set_run_font(r, 8.5, MUTED, bold=True)
    r = header.add_run(f"   |   {running_label}")
    set_run_font(r, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    r = footer.add_run("Tài liệu phát hành kèm bản build   •   Trang ")
    set_run_font(r, 8.5, MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_end)


def add_title(doc: Document, kicker: str, title: str, subtitle: str, detail: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(65)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(kicker.upper())
    set_run_font(r, 10, CYAN, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run_font(r, 29, NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(subtitle)
    set_run_font(r, 15, BLUE_DARK, bold=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run(detail)
    set_run_font(r, 10, MUTED, italic=True)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, 11, INK, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, 11, INK)
    else:
        r = p.add_run(text)
        set_run_font(r, 11, INK)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, 11, INK)


def add_step(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, 11, INK)


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_BLUE, accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    mark_header_row(table.rows[0])
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, color=accent, size="10")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper())
    set_run_font(r, 9.5, accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    mark_header_row(table.rows[0])
    dxa = [int(round(value * 1440)) for value in widths]
    set_table_geometry(table, dxa)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, 9.5, BLUE_DARK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, 9.5, INK)
    set_table_geometry(table, dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_diagram(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.45))
    inline = run._r.xpath(".//wp:docPr")
    if inline:
        inline[0].set("descr", caption)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = False


def add_picture_with_alt(paragraph, image_path: Path, width: float, alt_text: str) -> None:
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    inline = run._r.xpath(".//wp:docPr")
    if inline:
        inline[0].set("descr", alt_text)
        inline[0].set("title", alt_text)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def build_guide(output: Path, logo: Path, user_flow: Path, cycle: Path, media_flow: Path) -> None:
    doc = Document()
    configure_document(doc, "HƯỚNG DẪN NGƯỜI DÙNG")
    add_title(
        doc,
        "HƯỚNG DẪN NGƯỜI DÙNG",
        PRODUCT,
        "Bảo vệ dữ liệu cá nhân trên Windows bằng một cổng xác thực duy nhất",
        "Dành cho người dùng cuối • Không yêu cầu kiến thức kỹ thuật",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    add_picture_with_alt(p, logo, 1.18, "Biểu trưng FileSentry Sentinel hình khiên và dấu kiểm")
    add_callout(doc, "Bắt đầu nhanh", "Mở app → nhập mật khẩu trong FileSentry → chọn đúng quyền cần dùng → thao tác thật diễn ra trực tiếp trên Windows. Phiên xác thực thao tác thông thường là 15 phút; tài nguyên nhạy cảm có thể yêu cầu xác thực riêng và có thời hạn hiển thị trong app.", fill="EEF8FC", accent=BLUE)
    add_table(doc, ["Thông tin", "Giá trị"], [["Kênh phát hành", "V1 Local Defensive Console + V2 groundwork"], ["Đối tượng", "Người dùng cuối trên Windows"], ["Mục tiêu", "Sử dụng các quyền bảo vệ an toàn"]], [1.875, 4.625])
    add_page_break(doc)

    doc.add_heading("1. FileSentry Sentinel dùng để làm gì?", level=1)
    add_body(doc, "FileSentry Sentinel là lớp kiểm soát cục bộ cho các tài nguyên nhạy cảm: khu vực/thư mục bảo vệ, camera, microphone, media và một số trạng thái kết nối. Ứng dụng yêu cầu xác thực trước, tạo phiên sử dụng có thời hạn, ghi nhận hoạt động và tự đưa quyền về trạng thái an toàn khi phiên hết hạn hoặc bị thay đổi ngoài ý muốn.")
    add_body(doc, "Bản phát hành này không tự gửi mật khẩu, file, log hoặc telemetry ra Internet. Đây là V1 local defensive console có nền tảng V2; Windows Service thin-client, Minifilter, ETW và browser extension vẫn là phần phát triển tiếp theo.")
    add_callout(doc, "Nguyên tắc dễ nhớ", "Mật khẩu chỉ nhập trong FileSentry. Sau khi xác nhận, bạn thao tác bình thường trên Windows.", fill="F4F6F9", accent=BLUE_DARK)
    doc.add_heading("2. Lần đầu mở ứng dụng", level=1)
    add_step(doc, "Mở FileSentry Sentinel bằng biểu tượng trên Desktop hoặc Start Menu.")
    add_step(doc, "Đăng nhập bằng tài khoản do người phụ trách cung cấp. Nếu đây là lần đầu sử dụng, hãy đổi mật khẩu ngay trong mục Cài đặt hệ thống.")
    add_step(doc, "Kiểm tra thanh trạng thái ở cuối màn hình: phiên đăng nhập hợp lệ, chính sách cục bộ đang hoạt động và thời gian còn lại của phiên.")
    add_step(doc, "Mở mục Hướng dẫn sử dụng bất cứ lúc nào. Mỗi màn hình đều có popup tour trỏ tới đúng nút hoặc bảng dữ liệu trên giao diện.")
    add_step(doc, "Ở bước thao tác, bấm đúng control được viền vàng; popup tự chuyển bước. Nếu bấm sai, popup nhắc lại. Ở bước chỉ đọc, bấm ĐÃ XEM — TIẾP THEO.")
    add_step(doc, "Popup hoạt động như bong bóng hướng dẫn: phần còn lại của cửa sổ được làm mờ, chỉ vùng đang được chỉ dẫn còn sáng để tránh bấm nhầm.")
    add_callout(doc, "Bảo mật tài khoản", "Không gửi mật khẩu quản trị qua chat/email. Nếu nghi ngờ mật khẩu bị lộ, khóa tài nguyên và đổi mật khẩu ngay.", fill="FFF8E8", accent=AMBER)

    doc.add_heading("3. Luồng sử dụng chung", level=1)
    add_body(doc, "Mọi chức năng quan trọng đều đi qua cùng một cổng xác thực. Phiên xác thực thao tác thông thường kéo dài 15 phút. Các hành động nhạy cảm như Vault, Camera/Microphone, khôi phục hoặc gỡ ứng dụng có thể yêu cầu xác thực riêng; thời hạn grant tài nguyên được hiển thị trong chính ứng dụng.")
    add_body(doc, "Bạn có thể bấm chuột phải ở từng màn hình để mở menu thao tác theo ngữ cảnh. Menu chung luôn có Làm mới màn hình và Mở hướng dẫn; khi bấm trên bảng dữ liệu, menu chỉ hiện các thao tác phù hợp với mục đang chọn.")
    add_bullet(doc, "Media Library: xem/mở media, khóa xóa, chống gửi ra ngoài hoặc gỡ bảo vệ.")
    add_bullet(doc, "Kho mã hóa và Cách ly: khôi phục mục đang chọn.")
    add_bullet(doc, "Khu vực bảo vệ: xóa phạm vi, mở kho, gỡ quản lý hoặc xử lý Folder Lock.")
    add_bullet(doc, "Kết nối mạng và Persistence: quét lại bảng dữ liệu.")
    add_diagram(doc, user_flow, "Sơ đồ: từ mở FileSentry đến thao tác trực tiếp trên Windows.")
    doc.add_heading("3A. Chế độ giao diện", level=2)
    add_body(doc, "Bấm biểu tượng giao diện ◐ / ☼ / ☾ ở góc phải header để mở menu. Bạn có thể chọn Ban ngày, Ban đêm hoặc Theo Windows; Theo Windows tự đọc thiết lập sáng/tối của Windows và tự cập nhật khi hệ thống thay đổi.")
    add_step(doc, "Từ bất kỳ màn hình nào, bấm biểu tượng giao diện ở góc phải header.")
    add_step(doc, "Chọn BAN NGÀY nếu muốn nền sáng, BAN ĐÊM nếu muốn nền tối hoặc THEO WINDOWS để đồng bộ với Windows.")
    add_body(doc, "Lựa chọn giao diện chỉ ảnh hưởng phần hiển thị; không thay đổi mật khẩu, dữ liệu, audit, chính sách bảo vệ hoặc trạng thái khóa tài nguyên. Cài đặt hệ thống chỉ hiển thị chế độ hiện tại.")
    doc.add_heading("4. Bảo vệ thư mục / khu vực file", level=1)
    add_body(doc, "FileSentry dùng một luồng thống nhất: Tổng quan → Phạm vi → Chính sách → Theo dõi → Xử lý. Thanh Protection Journey ở đầu mỗi màn hình cho phép chuyển giữa các bước; Access Center là điểm chọn tập trung cho Camera/Microphone, Media Library, Kho mã hóa và Khu vực bảo vệ.")
    add_body(doc, "Include/exclude là phạm vi để FileSentry giám sát event file; việc thêm một khu vực không tự chặn truy cập Explorer. Muốn khóa quyền Windows cho một thư mục cụ thể, dùng Folder Lock bên dưới. Chặn I/O toàn hệ thống cần Windows Service + Minifilter của V2.")
    add_step(doc, "Mở Khu vực bảo vệ ở thanh điều hướng.")
    add_step(doc, "Chọn Thêm khu vực, chọn thư mục cần bảo vệ và đặt tên dễ nhận biết.")
    add_step(doc, "Nếu cần một thư mục riêng để lưu trữ, chọn + Tạo kho lưu trữ, chọn thư mục gốc và nhập tên. FileSentry tạo thư mục, thêm vào phạm vi giám sát và giữ nguyên dữ liệu đang có.")
    add_step(doc, "Chọn kho rồi bấm Mở thư mục để mở bằng Windows Explorer. Gỡ quản lý (giữ thư mục) chỉ bỏ nhãn riêng trong app, không xóa thư mục, file hoặc phạm vi giám sát.")
    add_step(doc, "Bật bảo vệ trong Tổng quan sau khi đã chọn include/exclude. FileSentry bắt đầu giám sát theo phạm vi; không tự di chuyển hoặc xóa file.")
    add_step(doc, "Nếu cần mở thư mục bảo vệ, xác thực trong FileSentry rồi bấm Mở thư mục. Quyền truy cập thực tế của Explorer vẫn do Windows và Folder Lock quyết định.")
    add_step(doc, "Khi không cần theo dõi, dùng Tắt bảo vệ hoặc Tạm dừng theo thời lượng. Khi cần khóa quyền quản lý phạm vi, dùng Khóa khu vực; đây là khóa ở tầng ứng dụng MVP.")
    add_diagram(doc, cycle, "Sơ đồ: bốn trạng thái người dùng thường gặp của khu vực bảo vệ.")
    add_table(doc, ["Trạng thái", "Ý nghĩa", "Bạn nên làm gì"], [["Đang khóa", "Không cho truy cập theo chính sách hiện tại", "Xác thực rồi mở khóa khi thật sự cần"], ["Đang mở", "Phiên đã được cấp quyền", "Hoàn tất công việc rồi khóa lại"], ["Tạm dừng", "Cho phép trong thời gian ngắn", "Dùng cho tác vụ ngắn; kiểm tra thời gian còn lại"], ["Cảnh báo", "Phát hiện thay đổi hoặc lỗi", "Đọc thông báo, khóa lại và xem Nhật ký hoạt động"]], [1.35, 2.7, 2.45])
    doc.add_heading("4A. Khóa thư mục bằng Windows ACL", level=1)
    add_body(doc, "Folder Lock dùng quyền NTFS cho một thư mục cụ thể. Đây là khóa quyền, không phải mã hóa; Windows Administrator/SYSTEM vẫn là ranh giới của hệ điều hành.")
    add_step(doc, "Trong Khu vực bảo vệ, chọn + Khóa thư mục và xác thực trong FileSentry.")
    add_step(doc, "FileSentry lưu DACL gốc đã mã hóa trước khi áp khóa. Nếu không lưu được bản sao, khóa không được áp dụng.")
    add_step(doc, "Khi cần dùng lại, chọn mục và bấm Mở khóa mục chọn để khôi phục đúng DACL gốc.")
    add_step(doc, "Bấm Kiểm tra ACL nếu nghi ngờ quyền bị thay đổi từ bên ngoài. FileSentry chỉ cảnh báo, không tự sửa ngoài quy trình.")

    doc.add_heading("5. Bật, tắt và tạm dừng", level=1)
    add_body(doc, "Các nút Bật/Tắt bảo vệ, Tạm dừng và Khóa khu vực chỉ thay đổi trạng thái giám sát hoặc quyền quản lý trong FileSentry. Thao tác được bảo vệ bởi PasswordGate, có thể dùng lại phiên 15 phút cho việc thông thường và ghi audit; hành động rủi ro cao vẫn yêu cầu xác thực mới.")
    add_step(doc, "Từ Tổng quan, chọn Bật bảo vệ để tiếp tục theo dõi theo include/exclude hiện tại.")
    add_step(doc, "Chọn Tắt bảo vệ khi muốn dừng giám sát. Dữ liệu, cấu hình, log và Folder Lock không tự bị xóa.")
    add_step(doc, "Chọn Tạm dừng 15 phút khi cần thao tác hàng loạt; sau thời hạn, giám sát trở lại theo cấu hình.")
    add_step(doc, "Chọn Khóa khu vực để khóa màn hình quản lý include/exclude trong FileSentry. Đây chưa phải filesystem lock của Windows.")

    add_page_break(doc)
    doc.add_heading("6. Quản lý ảnh, video và âm thanh", level=1)
    add_body(doc, "Mục Ảnh / Video / Âm thanh giúp bạn tập trung các file media trên các ổ đĩa cục bộ, khóa xóa từng file và đưa dữ liệu riêng vào Kho riêng mã hóa.")
    add_step(doc, "Bấm Đồng bộ toàn bộ máy để kiểm tra các ổ đĩa cục bộ. File mới được thêm, file thay đổi được cập nhật, còn file đã xóa hoặc di chuyển được đánh dấu Đã rời khỏi máy.")
    add_step(doc, "Nếu chỉ muốn quản lý một file, bấm Thêm file media và chọn file đó.")
    add_step(doc, "Cửa sổ tiến trình hiển thị số file đã xử lý, phần trăm và thời gian ước tính. Có thể bấm Hủy đồng bộ; các file chưa kiểm tra sẽ không bị đánh dấu nhầm.")
    add_step(doc, "Sau lần đồng bộ đầu tiên, các thay đổi media trên những ổ đĩa đã quét được cập nhật tự động khi ứng dụng đang chạy. Khi gắn ổ đĩa mới, hãy đồng bộ toàn bộ máy lại.")
    add_step(doc, "Chọn Khóa xóa file nếu không muốn file bị xóa hoặc đổi tên từ Windows.")
    add_step(doc, "Chọn Chống gửi ra ngoài để đưa file vào Kho riêng mã hóa. Sau khi mã hóa thành công, bản file thường bên ngoài sẽ được xóa.")
    add_step(doc, "Nếu cần trả lại quyền thông thường, chọn Gỡ bảo vệ file và xác nhận.")
    add_step(doc, "Chọn một mục rồi bấm Xem / Mở media đã chọn. Ảnh được xem trong FileSentry có thanh cuộn; video và âm thanh file ngoài được mở bằng ứng dụng mặc định của Windows. Media trong Kho riêng không tạo file thường bên ngoài.")
    add_step(doc, "Bấm Xóa sạch danh sách (giữ file) nếu chỉ muốn xóa inventory file ngoài khỏi app. Nội dung file thật không bị xóa; ACL do FileSentry tạo được dọn, còn mục Kho riêng được giữ lại.")
    add_callout(doc, "Giới hạn cần nhớ", "Một file vẫn là file thường ngoài FileSentry thì không thể bị chặn tuyệt đối việc sao chép hoặc tải lên từ mọi ứng dụng. Muốn bảo vệ mạnh nhất, hãy dùng Kho riêng.", fill="FFF8E8", accent=AMBER)

    doc.add_heading("7. Camera và microphone", level=1)
    add_body(doc, "Trong mục Camera & Microphone, bạn có thể khóa, mở khóa, khóa tạm thời và lưu danh sách website origin được phép. Danh sách origin hiện là dữ liệu chính sách cho browser extension tương lai; chưa phải enforcement URL bên trong trình duyệt.")
    add_diagram(doc, media_flow, "Sơ đồ: xác thực một lần rồi dùng camera/microphone trực tiếp trên Windows hoặc trình duyệt.")
    add_step(doc, "Mở Camera & Microphone và chọn Camera hoặc Microphone.")
    add_step(doc, "Bấm Mở khóa, nhập mật khẩu trong FileSentry và xác nhận. Sau đó bạn sử dụng camera hoặc microphone bình thường.")
    add_step(doc, "Nếu chỉ cần dùng ngắn, chọn Khóa tạm thời và đặt thời lượng. Khi hết thời gian, chính sách trở về khóa.")
    add_step(doc, "Để chặn hoàn toàn, bấm Khóa. Nếu có website cụ thể, thêm đúng địa chỉ website đó vào danh sách được phép.")
    add_step(doc, "Nếu trạng thái không đúng mong muốn, đọc thông báo trong app và kiểm tra lại mục này trước khi sử dụng.")
    add_callout(doc, "Ghi nhớ", "Camera/Microphone hiện dùng Windows policy + watch-and-revert. Windows không cung cấp pre-hook tuyệt đối cho ứng dụng bên thứ ba; có thể có khoảng thời gian ngắn trước khi trạng thái bị đưa về khóa.", fill="FFF8E8", accent=AMBER)

    doc.add_heading("8. Mạng và cảnh báo kết nối", level=1)
    add_body(doc, "Mục Kết nối mạng chỉ đọc bảng socket TCP/UDP cục bộ, process/PID, địa chỉ local/remote và chỉ báo kết nối Internet. FileSentry không port-scan, không phân giải DNS, không upload telemetry và không tự động chặn network trong bản này.")
    add_bullet(doc, "Xem tên mạng, trạng thái kết nối và thời điểm kiểm tra gần nhất.")
    add_bullet(doc, "Mở Nhật ký hoạt động để xem sự kiện, tài nguyên liên quan và kết quả xử lý.")
    add_bullet(doc, "Nếu cảnh báo lặp lại: khóa các khu vực nhạy cảm, ngắt kết nối không cần thiết và liên hệ người phụ trách hệ thống.")
    add_body(doc, "FileSentry đối chiếu thay đổi file đáng chú ý với kết nối Internet bên ngoài trong cùng một cửa sổ thời gian. Khi có đủ bằng chứng, thẻ DOUBLE-EXTORTION CORRELATION xuất hiện ở Nhật ký hoạt động và Kết nối mạng; thẻ hiển thị số lượng, thời gian, endpoint và đường dẫn để bạn kiểm tra.")
    add_step(doc, "Mở Nhật ký hoạt động, đọc thẻ tương quan và bấm Mở bảng Network để kiểm tra process, endpoint và indicator.")
    add_step(doc, "Nếu cần cô lập file, bấm Mở cách ly và chọn file cụ thể. Tương quan là chỉ báo cần kiểm tra, không phải kết luận máy đã bị xâm nhập.")
    add_callout(doc, "An toàn", "Không mở quyền cho camera, microphone hoặc thư mục khi bạn chưa chủ động xác nhận trong FileSentry.", fill="EEF8FC", accent=BLUE)

    doc.add_heading("9. Cách ly, Kho mã hóa và khôi phục", level=1)
    add_body(doc, "Cách ly đưa file nghi ngờ vào vùng lưu trữ mã hóa; Kho mã hóa lưu bản mã hóa theo từng file/chunk. Đây là lưu trữ và khôi phục có kiểm tra hash, chưa phải khóa Explorer real-time.")
    add_step(doc, "Trong Cách ly, chọn file rồi bấm Khôi phục mục đang chọn. FileSentry yêu cầu xác thực và không tự ghi đè file đích đang tồn tại.")
    add_step(doc, "Trong Kho mã hóa, bấm Đưa file vào Vault, xác thực và chọn file. Bản gốc được giữ theo chính sách V1; bản trong Vault được mã hóa.")
    add_step(doc, "Chọn mục Vault và bấm Khôi phục mục đang chọn để chọn đường dẫn đích. FileSentry kiểm tra hash trước khi hoàn tất.")
    add_step(doc, "Trong Nhật ký hoạt động, bấm Xuất báo cáo sự cố, chọn khoảng thời gian và nơi lưu file .fsreport. Báo cáo được mã hóa cục bộ.")
    add_callout(doc, "Không tự kết luận", "Cảnh báo, tương quan file/network và báo cáo chỉ là bằng chứng cần kiểm tra; không phải kết luận chắc chắn máy đã bị xâm nhập.", fill="FFF8E8", accent=AMBER)

    doc.add_heading("10. Persistence, sao lưu và gỡ ứng dụng", level=1)
    add_step(doc, "Mở Startup & Persistence để xem Run/RunOnce, Startup folder, Scheduled Task và Windows Service. FileSentry chỉ inventory/read-only, không tự xóa entry.")
    add_step(doc, "Sao lưu thư mục dữ liệu FileSentry theo quy trình nội bộ và kiểm thử khôi phục định kỳ; không gửi mật khẩu hoặc file riêng tư ra ngoài.")
    add_step(doc, "Muốn gỡ ứng dụng, vào Cài đặt hệ thống → Gỡ FileSentry, mở khóa và kiểm tra tất cả Folder Lock trước, sau đó xác nhận từng bước. Dữ liệu chỉ bị xóa khi bạn chọn xóa toàn bộ.")
    add_callout(doc, "Giới hạn phát hành", "Service lifecycle/recovery, thin-client, ETW/Sysmon, Minifilter, browser extension và external hash anchor chưa phải tính năng hoàn chỉnh của bản này.", fill="EEF8FC", accent=BLUE)

    doc.add_heading("11. Thông báo và cách đọc", level=1)
    add_table(doc, ["Loại", "Bạn sẽ thấy", "Hành động đề xuất"], [["Thành công", "Đã mở khóa / đã khóa / đã lưu", "Tiếp tục hoặc khóa lại sau khi xong"], ["Thông tin", "Phiên còn bao lâu, chính sách nào đang hoạt động", "Đọc để biết trạng thái hiện tại"], ["Cảnh báo", "Thay đổi ngoài phiên, kết nối đáng chú ý", "Không bỏ qua; kiểm tra Nhật ký"], ["Lỗi", "Không thể áp dụng chính sách hoặc mất kết nối", "Giữ tài nguyên ở trạng thái khóa và thử lại sau"]], [1.25, 2.75, 2.5])
    doc.add_heading("12. Hướng dẫn xử lý nhanh", level=1)
    add_table(doc, ["Tình huống", "Cách xử lý"], [["Quên mật khẩu", "Không đoán liên tục. Liên hệ người sở hữu/quản trị để xử lý theo quy trình khôi phục."], ["Đã mở khóa nhưng không truy cập được", "Kiểm tra phiên còn hiệu lực, tên khu vực, quyền riêng tư của Windows và Nhật ký hoạt động."], ["Camera/mic vẫn không hoạt động", "Kiểm tra thiết bị vật lý, trình duyệt và quyền riêng tư của Windows."], ["Cảnh báo xuất hiện liên tục", "Khóa tài nguyên, ghi lại thời điểm và liên hệ người phụ trách."], ["Xuất hiện cửa sổ hệ thống", "Không nhập nội dung lạ. Đóng cửa sổ nếu có thể và xem lại thông báo trong FileSentry."]], [2.1, 4.4])

    doc.add_heading("13. Quy tắc sử dụng an toàn", level=1)
    for text in [
        "Luôn khóa lại khu vực, camera và microphone khi không sử dụng.",
        "Chỉ thêm website bạn nhận diện rõ; không thêm địa chỉ lạ để “thử”.",
        "Không tắt FileSentry khi còn tài nguyên đang mở hoặc đang có cảnh báo.",
        "Giữ bản hướng dẫn và chứng chỉ build đi kèm cùng một thư mục phát hành.",
        "Nếu cần thay đổi chính sách, thực hiện trong app và đọc popup xác nhận trước khi bấm tiếp tục.",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "Tóm tắt một câu", "FileSentry là ổ khóa và cổng xác thực; Windows là nơi thao tác thật. Hãy xác thực trong app, dùng đúng thời gian cần thiết, rồi khóa lại.", fill="EEF8FC", accent=GREEN)
    doc.add_heading("Khi cần hỗ trợ", level=1)
    add_body(doc, "Hãy ghi lại màn hình đang mở, thời điểm xảy ra và nội dung thông báo. Không gửi mật khẩu, mã xác thực hoặc tệp dữ liệu riêng tư khi yêu cầu hỗ trợ.")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_certificate(output: Path, owner: str, built_at: str, logo: Path, exe_sha256: str) -> None:
    doc = Document()
    configure_document(doc, "CHỨNG CHỈ BUILD NỘI BỘ")
    add_title(
        doc,
        "CHỨNG CHỈ BUILD NỘI BỘ",
        "BẢN BUILD ĐỘC QUYỀN",
        PRODUCT,
        "Tài liệu xác nhận quyền sở hữu nội bộ của bản phát hành",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    add_picture_with_alt(p, logo, 1.0, "Biểu trưng FileSentry Sentinel hình khiên và dấu kiểm")
    add_callout(doc, "Trạng thái", "Tài liệu này được tạo cùng bản phát hành nội bộ. Hãy giữ nguyên và chỉ chia sẻ cho người được chủ sở hữu cho phép.", fill="EEF8FC", accent=BLUE)
    add_table(doc, ["Thuộc tính", "Giá trị"], [["Sản phẩm", PRODUCT], ["Chủ sở hữu", owner], ["Ngày phát hành", built_at.split(" UTC")[0]], ["Kênh phát hành", "V1 Local Defensive Console + V2 groundwork"], ["EXE SHA-256", exe_sha256], ["Mục đích", "Xác nhận quyền sở hữu và đối chiếu đúng file EXE"]], [1.875, 4.625])
    doc.add_heading("Ý nghĩa chứng chỉ", level=1)
    add_body(doc, "Chứng chỉ xác nhận sản phẩm được phát hành dưới quyền sở hữu của người đứng tên trong tài liệu tại thời điểm phát hành. EXE SHA-256 là dấu vân tay để đối chiếu bản EXE đi kèm; nếu hash khác, không coi đó là cùng một bản phát hành.")
    add_callout(doc, "Giới hạn xác minh", "Word read-only chỉ ngăn chỉnh sửa thông thường trong giao diện Word, không phải mã hóa hay chữ ký số tuyệt đối. Để xác minh, đối chiếu EXE SHA-256 và chủ sở hữu phát hành.", fill="FFF8E8", accent=AMBER)
    doc.add_heading("Cách sử dụng", level=1)
    add_step(doc, "Giữ chứng chỉ cùng bộ tài liệu phát hành.")
    add_step(doc, "Không chỉnh sửa, sao chép nội dung sang tài liệu khác hoặc chia sẻ cho người không được phép.")
    add_step(doc, "Nếu nội dung hoặc nguồn phát hành có dấu hiệu bất thường, liên hệ trực tiếp với chủ sở hữu.")
    add_callout(doc, "Phạm vi sử dụng", "Đây là tài liệu sở hữu nội bộ của FileSentry Sentinel, dùng cho mục đích nhận diện và quản lý bản phát hành.", fill="FFF8E8", accent=AMBER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("FILESENTRY SENTINEL • TÀI LIỆU SỞ HỮU NỘI BỘ")
    set_run_font(r, 9, MUTED, bold=True)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--exe", required=True, type=Path)
    parser.add_argument("--dist", required=True, type=Path)
    parser.add_argument("--owner", default="Local Project Owner")
    parser.add_argument("--guide-source", type=Path, help="Kept for build-script compatibility; the DOCX guide is authored here.")
    parser.add_argument("--certificate-output", type=Path, help="Optional draft certificate output path.")
    args = parser.parse_args()
    if not args.exe.is_file():
        raise SystemExit(f"EXE not found: {args.exe}")
    if args.guide_source and not args.guide_source.is_file():
        raise SystemExit(f"Guide source not found: {args.guide_source}")

    built_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    exe_sha256 = sha256_file(args.exe)
    guide = args.dist / "FileSentrySentinel_User_Guide.docx"
    certificate = args.certificate_output or (args.dist / "FileSentrySentinel_Exclusive_Build_Certificate.docx")
    with tempfile.TemporaryDirectory(prefix="filesentry-docs-") as temp_dir:
        temp = Path(temp_dir)
        logo = temp / "filesentry-logo.png"
        user_flow = temp / "user-flow.png"
        cycle = temp / "protection-cycle.png"
        media_flow = temp / "media-flow.png"
        make_logo(logo)
        make_user_flow(user_flow)
        make_protection_cycle(cycle)
        make_media_flow(media_flow)
        build_guide(guide, logo, user_flow, cycle, media_flow)
        build_certificate(certificate, args.owner, built_at, logo, exe_sha256)
    print(f"[OK] wrote {guide}")
    print(f"[OK] wrote {certificate}")


if __name__ == "__main__":
    main()
